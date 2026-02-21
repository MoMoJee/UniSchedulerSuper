"""
文档解析器

支持：
- 百度云智能文档解析（云端，异步提交任务 + 转为 Markdown）
- PDF (pdfplumber 本地局部备用)
- Word (python-docx 本地局部备用)
- Excel (openpyxl 本地局部备用)

降级链：baidu（云端） → pdfplumber/python-docx/openpyxl（本地）
所有本地解析器优雅降级：如果依赖库未安装，返回友好提示而不是崩溃。
"""
from typing import Dict, Any
import os
import time
import base64
import urllib.parse

from .base import BaseParser
from logger import logger


def _table_to_markdown(table_data) -> str:
    """将二维数组转换为 Markdown 表格"""
    if not table_data or len(table_data) == 0:
        return ""
    
    lines = []
    header = table_data[0]
    col_count = len(header)
    
    lines.append('| ' + ' | '.join(str(cell or '') for cell in header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * col_count) + ' |')
    
    for row in table_data[1:]:
        # 补齐列数
        padded = list(row) + [''] * (col_count - len(row))
        lines.append('| ' + ' | '.join(str(cell or '') for cell in padded[:col_count]) + ' |')
    
    return '\n'.join(lines)


class PDFParser(BaseParser):
    """PDF 解析器"""
    
    def can_parse(self, mime_type: str) -> bool:
        return mime_type == 'application/pdf'
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        try:
            import pdfplumber
        except ImportError:
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": "需要安装 pdfplumber 库: pip install pdfplumber"
            }
        
        try:
            text_content = []
            metadata = {"pages": 0, "has_tables": False}
            
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    page_texts = []
                    
                    # 提取文字
                    page_text = page.extract_text()
                    if page_text:
                        page_texts.append(page_text)
                    
                    # 提取表格
                    tables = page.extract_tables()
                    if tables:
                        metadata["has_tables"] = True
                        for table in tables:
                            if table:
                                table_md = _table_to_markdown(table)
                                if table_md:
                                    page_texts.append(table_md)
                    
                    if page_texts:
                        text_content.append(f"--- 第 {i+1} 页 ---\n" + '\n\n'.join(page_texts))
            
            full_text = '\n\n'.join(text_content)
            
            return {
                "success": True,
                "text": full_text or "[PDF 无可提取的文本内容]",
                "metadata": metadata,
                "error": ""
            }
            
        except Exception as e:
            logger.error(f"PDF 解析失败 {file_path}: {e}")
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": str(e)
            }


class WordParser(BaseParser):
    """Word 文档解析器"""
    
    SUPPORTED_MIMES = {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
    }
    
    def can_parse(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_MIMES
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError:
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": "需要安装 python-docx 库: pip install python-docx"
            }
        
        try:
            doc = Document(file_path)
            
            # 提取段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 提取表格
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables_text.append(_table_to_markdown(table_data))
            
            full_text = '\n\n'.join(paragraphs + tables_text)
            
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
            }
            
            return {
                "success": True,
                "text": full_text or "[Word 文档无内容]",
                "metadata": metadata,
                "error": ""
            }
            
        except Exception as e:
            logger.error(f"Word 解析失败 {file_path}: {e}")
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": str(e)
            }


class ExcelParser(BaseParser):
    """Excel 表格解析器"""
    
    SUPPORTED_MIMES = {
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel',
    }
    
    def can_parse(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_MIMES
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        try:
            import openpyxl
        except ImportError:
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": "需要安装 openpyxl 库: pip install openpyxl"
            }
        
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheets_text = []
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                
                data = []
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append([
                            str(cell) if cell is not None else '' 
                            for cell in row
                        ])
                
                if data:
                    sheet_md = f"### Sheet: {sheet_name}\n\n"
                    sheet_md += _table_to_markdown(data)
                    sheets_text.append(sheet_md)
            
            full_text = '\n\n'.join(sheets_text)
            
            metadata = {
                "sheets": len(wb.sheetnames),
                "sheet_names": wb.sheetnames,
            }
            
            return {
                "success": True,
                "text": full_text or "[Excel 文件无数据]",
                "metadata": metadata,
                "error": ""
            }
            
        except Exception as e:
            logger.error(f"Excel 解析失败 {file_path}: {e}")
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": str(e)
            }


# ============================================================
# 百度云智能文档解析器（云端优先，全格式局部降级）
# ============================================================

class BaiduDocumentParser(BaseParser):
    """
    百度云智能文档解析器

    支持格式：PDF / Word (.docx/.doc) / Excel (.xlsx/.xls)

    流程：
      1. 透过 AK/SK 获取（或复用）access_token
      2. 提交文档任务，得到 task_id
      3. 轮询任务状态，直到 success / failed / 超时
      4. 下载 markdown_url 得到文本

    降级：任何阶段失败时自动调用对应的本地解析器。
    配置確认：`document_services.baidu.enabled = true` 时才尝试云端。
    """

    SUPPORTED_MIMES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel',
    }

    # 类层面 access_token 缓存（所有实例共享）
    _access_token: str = ""
    _token_expires_at: float = 0.0

    def can_parse(self, mime_type: str) -> bool:
        """仅在云端服务已配置并启用时返回 True"""
        if mime_type not in self.SUPPORTED_MIMES:
            return False
        # 惰性检查配置（不加载，只看缓存）——工厂初始化时调用，配置应已加载
        try:
            from config.api_keys_manager import APIKeyManager
            return APIKeyManager.get_document_service_config('baidu') is not None
        except Exception:
            return False

    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        解析文档：云端为主，本地为备。

        Returns:
            {
                "success": bool,
                "text": str,           # Markdown 文本
                "metadata": dict,
                "error": str,
                "source": str,         # 'baidu_cloud' | 'local_fallback'
            }
        """
        mime_type = kwargs.get('mime_type', '')
        filename = os.path.basename(file_path)

        # 尝试云端解析
        logger.info(f"[文档解析-百度云] 开始: {filename}")
        cloud_result = self._try_baidu_cloud(file_path, filename)
        if cloud_result.get('success'):
            cloud_result['source'] = 'baidu_cloud'
            return cloud_result

        # 云端失败，降级本地
        logger.warning(
            f"[文档解析] 百度云端失败，回退本地解析: {filename}. "
            f"云端错误: {cloud_result.get('error', '')}"
        )
        local_result = self._local_fallback(file_path, mime_type, filename)
        local_result['source'] = 'local_fallback'
        return local_result

    # ---- 百度云主流程 ----

    def _try_baidu_cloud(self, file_path: str, filename: str) -> Dict[str, Any]:
        """complete 百度云文档解析流程：提交 → 轮询 → 下载"""
        try:
            import requests
            from config.api_keys_manager import APIKeyManager
        except ImportError as e:
            return {"success": False, "text": "", "metadata": {}, "error": f"缺少依赖: {e}"}

        config = APIKeyManager.get_document_service_config('baidu')
        if not config:
            return {"success": False, "text": "", "metadata": {},
                    "error": "百度云文档解析权数未配置或已禁用"}

        # 1. 获取 access_token
        token = self._get_access_token(config)
        if not token:
            return {"success": False, "text": "", "metadata": {},
                    "error": "获取 access_token 失败，请检查 api_key/secret_key"}

        # 2. 读取文件并 base64 编码
        try:
            with open(file_path, 'rb') as f:
                file_b64 = base64.b64encode(f.read()).decode('utf-8')
        except Exception as e:
            return {"success": False, "text": "", "metadata": {}, "error": f"文件读取失败: {e}"}

        # 3. 提交解析任务
        submit_url = config.get('submit_url',
            'https://aip.baidubce.com/rest/2.0/brain/online/v2/parser/task')
        try:
            # 与测试代码保持一致：file_data=<urlencoded_base64>
            # 额外加 file_name，让百度服务端能识别文件格式（PDF 必须）
            payload = (
                'file_data=' + urllib.parse.quote_plus(file_b64)
                + '&file_name=' + urllib.parse.quote_plus(filename)
            )
            resp = requests.post(
                f"{submit_url}?access_token={token}",
                headers={'Content-Type': 'application/x-www-form-urlencoded',
                         'Accept': 'application/json'},
                data=payload.encode('utf-8'),
                timeout=60,
            )
            resp.encoding = 'utf-8'
            submit_data = resp.json()
        except requests.exceptions.Timeout:
            return {"success": False, "text": "", "metadata": {}, "error": "提交任务请求超时"}
        except Exception as e:
            return {"success": False, "text": "", "metadata": {}, "error": f"提交任务失败: {e}"}

        if submit_data.get('error_code', 0) != 0:
            err = f"提交失败 (code={submit_data.get('error_code')}): {submit_data.get('error_msg', '')}"
            logger.warning(f"[文档解析-百度云] {err} {filename}")
            if submit_data.get('error_code') in (110, 111):
                BaiduDocumentParser._access_token = ''
                BaiduDocumentParser._token_expires_at = 0.0
            return {"success": False, "text": "", "metadata": {}, "error": err}

        task_id = submit_data.get('result', {}).get('task_id', '')
        if not task_id:
            return {"success": False, "text": "", "metadata": {},
                    "error": f"返回的任务 ID 为空: {submit_data}"}

        logger.info(f"[文档解析-百度云] 任务已提交: {task_id}, 文件: {filename}")

        # 4. 轮询任务状态
        query_url = config.get('query_url',
            'https://aip.baidubce.com/rest/2.0/brain/online/v2/parser/task/query')
        interval = float(config.get('poll_interval_seconds', 2))
        max_attempts = int(config.get('poll_max_attempts', 30))
        markdown_url = None

        for attempt in range(1, max_attempts + 1):
            time.sleep(interval)
            try:
                poll_resp = requests.post(
                    f"{query_url}?access_token={token}",
                    headers={'Content-Type': 'application/x-www-form-urlencoded',
                             'Accept': 'application/json'},
                    data=f'task_id={task_id}'.encode('utf-8'),
                    timeout=15,
                )
                poll_resp.encoding = 'utf-8'
                poll_data = poll_resp.json()
            except Exception as e:
                logger.warning(f"[文档解析-百度云] 轮询异常 (attempt {attempt}): {e}")
                continue

            if poll_data.get('error_code', 0) != 0:
                err = f"轮询失败 (code={poll_data.get('error_code')}): {poll_data.get('error_msg', '')}"
                logger.warning(f"[文档解析-百度云] {err}")
                return {"success": False, "text": "", "metadata": {}, "error": err}

            result = poll_data.get('result', {})
            status = result.get('status', '')
            logger.debug(f"[文档解析-百度云] 任务状态 ({attempt}/{max_attempts}): {status}")

            if status == 'success':
                markdown_url = result.get('markdown_url', '')
                break
            elif status == 'failed':
                task_err = result.get('task_error') or '未知错误'
                logger.warning(
                    f"[文档解析-百度云] 任务失败: {task_err}. "
                    "云端解析不可用，将尝试本地解析器。"
                )
                return {"success": False, "text": "", "metadata": {},
                        "error": f"百度云任务失败: {task_err}"}
            # 其他状态（pending/processing）继续等待
        else:
            # 超出最大轮询次数
            wait_secs = int(interval * max_attempts)
            logger.warning(
                f"[文档解析-百度云] 任务 {task_id} 超时（等待 {wait_secs}s）。"
                "云端解析超时，将尝试本地解析器。"
            )
            return {"success": False, "text": "", "metadata": {},
                    "error": f"百度云任务轮询超时（>{wait_secs}s）"}

        if not markdown_url:
            return {"success": False, "text": "", "metadata": {},
                    "error": "markdown_url 为空"}

        # 5. 下载 Markdown
        try:
            md_resp = requests.get(markdown_url, timeout=30)
            md_resp.encoding = 'utf-8'
            markdown_text = md_resp.text.strip()
        except Exception as e:
            err = f"下载 Markdown 失败: {e}"
            logger.warning(f"[文档解析-百度云] {err}，将尝试本地解析器。")
            return {"success": False, "text": "", "metadata": {}, "error": err}

        logger.info(
            f"[文档解析-百度云] 解析完成: {filename}, "
            f"{len(markdown_text)} 字符"
        )
        return {
            "success": True,
            "text": markdown_text or "[文档无可提取的内容]",
            "metadata": {
                "task_id": task_id,
                "markdown_url": markdown_url,
                "source": "baidu_cloud",
            },
            "error": "",
        }

    def _get_access_token(self, config: dict) -> str:
        """获取百度 access_token，带类级缓存（提前 5 分钟刷新）"""
        if (BaiduDocumentParser._access_token
                and time.time() < BaiduDocumentParser._token_expires_at - 300):
            return BaiduDocumentParser._access_token

        try:
            import requests
            token_url = config.get('token_url',
                'https://aip.baidubce.com/oauth/2.0/token')
            params = {
                'grant_type': 'client_credentials',
                'client_id': config['api_key'],
                'client_secret': config['secret_key'],
            }
            resp = requests.post(token_url, params=params, timeout=10)
            data = resp.json()
            token = data.get('access_token', '')
            expires_in = data.get('expires_in', 0)

            if not token:
                error = data.get('error_description', data.get('error', '未知错误'))
                logger.warning(
                    f"[文档解析-百度云] 获取 access_token 失败: {error}。"
                    "请检查 document_services.baidu.api_key / secret_key 是否正确。"
                )
                return ''

            BaiduDocumentParser._access_token = token
            BaiduDocumentParser._token_expires_at = time.time() + expires_in
            logger.info(
                f"[文档解析-百度云] access_token 获取成功，"
                f"有效期 {expires_in // 86400} 天"
            )
            return token
        except Exception as e:
            logger.warning(f"[文档解析-百度云] 获取 access_token 异常: {e}")
            return ''

    # ---- 本地降级 ----

    def _local_fallback(self, file_path: str, mime_type: str, filename: str) -> Dict[str, Any]:
        """根据 MIME 类型调用对应本地解析器"""
        pdf_mimes = {'application/pdf'}
        word_mimes = {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
        }
        excel_mimes = {
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
        }

        if mime_type in pdf_mimes:
            return PDFParser().parse(file_path)
        elif mime_type in word_mimes:
            return WordParser().parse(file_path)
        elif mime_type in excel_mimes:
            return ExcelParser().parse(file_path)
        else:
            # mime_type 可能为空（历史记录缺失等情况），按文件名后缀猜测
            ext = os.path.splitext(filename)[1].lower()
            ext_map = {
                '.pdf': PDFParser,
                '.docx': WordParser, '.doc': WordParser,
                '.xlsx': ExcelParser, '.xls': ExcelParser,
            }
            parser_cls = ext_map.get(ext)
            if parser_cls:
                return parser_cls().parse(file_path)
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": f"无法找到匹配的本地解析器 (mime={mime_type}, ext={ext})"
            }
